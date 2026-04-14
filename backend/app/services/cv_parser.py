"""
CV Parser — Pipeline Hybride Stratifié
=======================================
Architecture en 4 couches :

    Couche 1 — Regex         : email, téléphone, URLs (déterministe, haute précision)
    Couche 2 — NER (spaCy)   : noms propres, organisations, localisations
    Couche 3 — LLM (Groq)    : sémantique, compétences, résumé, expériences structurées
    Couche 4 — Post-processing: normalisation, validation croisée, déduplication

Formats supportés :
    - PDF (texte natif)
    - PDF multi-colonnes → DocLing (layout-aware, IBM Research 2024)
    - PDF image / CV scanné → OCR (Tesseract)
    - DOCX
    - TXT / Markdown
    - Image directe (JPG, PNG, WEBP, TIFF)

Dépendances :
    pip install groq pdfminer.six python-docx spacy pillow pytesseract python-dateutil langdetect
    pip install docling                                    # optionnel, pour les CVs multi-colonnes
    python -m spacy download fr_core_news_sm
    apt-get install tesseract-ocr tesseract-ocr-fra tesseract-ocr-eng
"""

from __future__ import annotations

import json
import logging
import os
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Optional

# ── Extraction de texte ───────────────────────────────────────────────────────
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.layout import LAParams, LTTextBox
from pdfminer.high_level import extract_pages
from docx import Document as DocxDocument

# ── OCR ───────────────────────────────────────────────────────────────────────
from PIL import Image
import pytesseract

# ── NER ───────────────────────────────────────────────────────────────────────
import spacy
from langdetect import detect

# ── LLM ───────────────────────────────────────────────────────────────────────
from groq import Groq
from dotenv import load_dotenv

# ── Normalisation des dates ───────────────────────────────────────────────────
from dateutil import parser as dateutil_parser

load_dotenv()


# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("cv_parser")


# ---------------------------------------------------------------------------
# Modèle de données
# ---------------------------------------------------------------------------

@dataclass
class CVData:
    """
    Structure de sortie normalisée du parser.
    
    """
    full_name:      Optional[str] = None
    email:          Optional[str] = None
    phone:          Optional[str] = None
    location:       Optional[str] = None
    linkedin:       Optional[str] = None
    github:         Optional[str] = None
    summary:        Optional[str] = None
    skills:         list[str]     = field(default_factory=list)
    languages:      list[dict]    = field(default_factory=list)  # [{language, level}]
    experiences:    list[dict]    = field(default_factory=list)
    education:      list[dict]    = field(default_factory=list)
    certifications: list[str]     = field(default_factory=list)

    # Métadonnées d'extraction (non-ATS, utiles pour debugging)
    _extraction_meta: dict = field(default_factory=dict)

    def to_json(self, indent: int = 2, include_meta: bool = False) -> str:
        d = asdict(self)
        if not include_meta:
            d.pop("_extraction_meta", None)
        return json.dumps(d, ensure_ascii=False, indent=indent)

    @classmethod
    def from_dict(cls, data: dict) -> "CVData":
        return cls(
            full_name      = data.get("full_name"),
            email          = data.get("email"),
            phone          = data.get("phone"),
            location       = data.get("location"),
            linkedin       = data.get("linkedin"),
            github         = data.get("github"),
            summary        = data.get("summary"),
            skills         = data.get("skills", []),
            languages      = data.get("languages", []),
            experiences    = data.get("experiences", []),
            education      = data.get("education", []),
            certifications = data.get("certifications", []),
        )


# ---------------------------------------------------------------------------
# COUCHE 0 — Extraction de texte brut (tous formats)
# ---------------------------------------------------------------------------

class TextExtractor:
    """
    Extrait le texte brut depuis n'importe quel format de CV.

    Stratégie PDF (par ordre de tentative) :
        1. pdfminer  — rapide, texte natif linéaire
        2. DocLing   — si pdfminer détecte une mise en page multi-colonnes
        3. Tesseract — si le PDF est scanné (texte < seuil minimal)

    """

    # Nombre de chars en dessous duquel le PDF est considéré comme scanné
    _PDF_MIN_CHARS = 50

    # Formats image supportés directement (OCR pur)
    _IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".tiff", ".tif", ".bmp"}

    @classmethod
    def extract(cls, file_path: str) -> tuple[str, str]:
        """
        Extrait le texte du fichier.

        Returns:
            (text, method) où method ∈ {'native', 'docling', 'ocr', 'docx', 'txt'}
        """
        path   = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return cls._extract_pdf(file_path)
        elif suffix == ".docx":
            return cls._extract_docx(file_path), "docx"
        elif suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8"), "txt"
        elif suffix in cls._IMAGE_EXTENSIONS:
            return cls._extract_image(file_path), "ocr"
        else:
            raise ValueError(
                f"Format non supporté : {suffix}. "
                f"Formats acceptés : pdf, docx, txt, md, jpg, png, webp, tiff"
            )

    # ── PDF ──────────────────────────────────────────────────────────────────

    @classmethod
    def _extract_pdf(cls, file_path: str) -> tuple[str, str]:
        """
        Orchestre l'extraction PDF selon la nature du document détectée.

        Ordre de tentative :
            1. pdfminer — extraction texte natif (rapide, gratuit)
            2. DocLing  — si pdfminer produit un texte bien extrait mais
                          que le layout spatial révèle des colonnes multiples
            3. Tesseract — si pdfminer n'a rien extrait (PDF scanné)
        """
        # Tentative 1 : pdfminer
        try:
            native_text = (pdf_extract_text(file_path) or "").strip()
        except Exception as e:
            logger.warning("pdfminer échoué : %s → bascule OCR", e)
            return cls._ocr_pdf(file_path), "ocr"

        # Cas : PDF scanné — pdfminer n'a pas extrait de texte significatif
        if len(native_text) < cls._PDF_MIN_CHARS:
            logger.info("PDF scanné détecté (%d chars) → OCR Tesseract", len(native_text))
            return cls._ocr_pdf(file_path), "ocr"

        # Cas : PDF texte natif — on vérifie si la mise en page est multi-colonnes
        if cls._has_multiple_columns(file_path):
            logger.info("PDF multi-colonnes détecté → DocLing")
            docling_text = cls._extract_pdf_docling(file_path)
            if docling_text:
                return docling_text, "docling"
            # DocLing non disponible ou échoué → on continue avec pdfminer
            logger.warning("DocLing indisponible ou échoué → fallback pdfminer")

        logger.info("PDF natif extrait via pdfminer (%d chars)", len(native_text))
        return native_text, "native"

    @classmethod
    def _has_multiple_columns(cls, file_path: str) -> bool:
        """
        Détecte si un PDF utilise une mise en page multi-colonnes.

        Méthode :
            On analyse la position horizontale (x0) de chaque bloc de texte
            (LTTextBox) sur la première page via les métadonnées spatiales
            que pdfminer expose nativement.

            Un CV deux colonnes présente deux groupes de blocs bien séparés
            sur l'axe horizontal. On mesure l'équilibre entre les blocs situés
            dans la moitié gauche et ceux dans la moitié droite de la page.
            Si chaque moitié contient au moins 25% des blocs, c'est multi-colonnes.

            Avantages de cette approche vs heuristique sur les lignes courtes :
              - Utilise les vraies coordonnées spatiales du PDF (fiable)
              - Ne confond pas une liste de skills (lignes courtes légitimes)
                avec un désordre de colonnes
              - Fonctionne dès la première page, pas besoin de tout lire

        Args:
            file_path: Chemin vers le fichier PDF.

        Returns:
            True si la mise en page est multi-colonnes, False sinon.
        """
        try:
            x_positions = []

            # On analyse uniquement la première page 
            for page_layout in extract_pages(file_path, laparams=LAParams(), page_numbers=[0]):
                for element in page_layout:
                    if isinstance(element, LTTextBox) and element.get_text().strip():
                        x_positions.append(element.x0)
                break  # première page uniquement

            if len(x_positions) < 4:
                # Pas assez de blocs pour conclure (page quasi-vide)
                return False

            # Séparer les blocs en deux moitiés horizontales
            x_min   = min(x_positions)
            x_max   = max(x_positions)
            midpoint = (x_min + x_max) / 2

            left_blocks  = sum(1 for x in x_positions if x <  midpoint)
            right_blocks = sum(1 for x in x_positions if x >= midpoint)
            total        = left_blocks + right_blocks

            # Équilibre gauche/droite — seuil à 25% pour chaque côté
            # Un CV une colonne a ~90% des blocs à gauche (alignement gauche)
            # Un CV deux colonnes a une distribution 40/60 à 60/40
            balance = min(left_blocks, right_blocks) / total
            is_multi = balance > 0.25

            logger.info(
                "Layout PDF : %d blocs (gauche=%d, droite=%d, balance=%.0f%%) → %s",
                total, left_blocks, right_blocks, balance * 100,
                "multi-colonnes" if is_multi else "une colonne",
            )
            return is_multi

        except Exception as e:
            logger.debug("Détection colonnes échouée : %s → one-column assumed", e)
            return False

    @classmethod
    def _extract_pdf_docling(cls, file_path: str) -> str:
        """
        Extrait le texte d'un PDF multi-colonnes avec DocLing (layout-aware).

        DocLing (IBM Research, 2024) analyse les coordonnées spatiales de chaque
        bloc et reconstruit l'ordre de lecture logique — colonnes, tableaux,
        en-têtes. C'est exactement ce dont on a besoin pour les CVs Canva/Word.

        Returns:
            Texte extrait par DocLing, ou "" si non disponible / erreur.
        """
        try:
            from docling.document_converter import DocumentConverter

            converter = DocumentConverter()
            result    = converter.convert(file_path)

            # DocLing expose le texte via export_to_text() ou export_to_markdown()
            # selon la version. On préfère le texte brut pour NER et LLM.
            if hasattr(result.document, "export_to_text"):
                return result.document.export_to_text()
            else:
                # Fallback markdown : on supprime les balises de formatage
                md = result.document.export_to_markdown()
                return re.sub(r"[#*`_>|-]+", " ", md)

        except ImportError:
            logger.debug(
                "DocLing non installé. "
                "Pour les CVs multi-colonnes : pip install docling"
            )
            return ""
        except Exception as e:
            logger.warning("DocLing échoué sur '%s' : %s", file_path, e)
            return ""

    @classmethod
    def _ocr_pdf(cls, file_path: str) -> str:
        """
        OCR d'un PDF scanné.
        Utilise pdf2image (meilleur rendu) si disponible, sinon PIL seul.
        """
        try:
            from pdf2image import convert_from_path  # type: ignore
            images = convert_from_path(file_path, dpi=300)
            logger.info("OCR PDF via pdf2image (%d pages)", len(images))
        except ImportError:
            logger.warning(
                "pdf2image non disponible — OCR dégradé. "
                "Installer : pip install pdf2image"
            )
            images = [Image.open(file_path)]

        pages_text = []
        for i, img in enumerate(images):
            text = pytesseract.image_to_string(
                img,
                lang="fra+eng",
                config="--psm 6",
            )
            pages_text.append(text)
            logger.debug("  Page %d OCR : %d chars", i + 1, len(text))

        return "\n\n".join(pages_text)

    # ── DOCX ─────────────────────────────────────────────────────────────────

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extraction DOCX : paragraphes + tableaux (souvent utilisés dans les CVs)."""
        doc   = DocxDocument(file_path)
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Certains CVs Word structurent les infos en grille
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    # ── Image ─────────────────────────────────────────────────────────────────

    @staticmethod
    def _extract_image(file_path: str) -> str:
        """OCR direct sur image (JPG, PNG, WEBP, etc.)."""
        img  = Image.open(file_path).convert("L")  # niveaux de gris
        text = pytesseract.image_to_string(img, lang="fra+eng", config="--psm 6")
        logger.info("OCR image : %d chars extraits", len(text))
        return text


# ---------------------------------------------------------------------------
# COUCHE 1 — Regex (champs déterministes)
# ---------------------------------------------------------------------------

class RegexExtractor:
    """
    Extraction haute-précision des champs à structure syntaxique fixe.
    Toujours exécutée EN PREMIER pour verrouiller les champs critiques.

    Principe : si le pattern matche, on a une confiance maximale.
    Le résultat Regex prime toujours sur NER et LLM pour ces champs.
    """

    # Email : RFC 5321 simplifié (couvre 99%+ des cas réels)
    EMAIL = re.compile(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
        re.IGNORECASE
    )

    # Téléphone — deux patterns complémentaires
    # PHONE_INTL : numéros avec préfixe international (+XX, 00XX, (+XX))
    PHONE_INTL = re.compile(
        r"(?:\(?\+\d{1,4}\)?|00\d{1,3})"
        r"[\s.\-]?"
        r"(?:\d[\s.\-]?)?"
        r"\d{2}(?:[\s.\-]?\d{2,3}){2,4}"
        r"(?!\d)"
    )
    # PHONE_LOCAL : numéros locaux FR/MA/TN (0X XX XX XX XX)
    PHONE_LOCAL = re.compile(
        r"(?<!\d)0[1-9](?:[\s.\-]?\d{2}){4}(?!\d)"
    )
    # Plages d'années — invalide les faux positifs (ex: "2020-2022")
    _YEAR_RANGE = re.compile(r"^(19|20)\d{2}[\s\-\u2013\u2014]+(19|20)\d{2}$")

    # URLs professionnelles
    LINKEDIN = re.compile(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/([\w\-\.]+)/?",
        re.IGNORECASE
    )
    GITHUB = re.compile(
        r"(?:https?://)?(?:www\.)?github\.com/([\w\-]+)/?",
        re.IGNORECASE
    )

    def extract(self, text: str) -> dict:
        """
        Extrait tous les champs déterministes du texte.

        Returns:
            dict avec clés : email, phone, linkedin, github (None si absent)
        """
        result = {}

        # Email
        email_match = self.EMAIL.search(text)
        result["email"] = email_match.group(0).lower() if email_match else None

        # Téléphone
        result["phone"] = self._extract_phone(text)

        # LinkedIn — URL normalisée
        linkedin_match = self.LINKEDIN.search(text)
        if linkedin_match:
            result["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"
        else:
            result["linkedin"] = None

        # GitHub — URL normalisée (exclusion des handles génériques)
        github_match = self.GITHUB.search(text)
        if github_match:
            handle = github_match.group(1)
            if handle.lower() not in ("features", "blog", "about", "contact", "login"):
                result["github"] = f"github.com/{handle}"
            else:
                result["github"] = None
        else:
            result["github"] = None

        logger.info(
            "Regex → email=%s | phone=%s | linkedin=%s | github=%s",
            bool(result["email"]), bool(result["phone"]),
            bool(result["linkedin"]), bool(result["github"]),
        )
        return result

    def _extract_phone(self, text: str) -> str | None:
        """
        Extrait le numéro de téléphone en combinant deux patterns.

        Stratégie : cherche d'abord les numéros internationaux (plus spécifiques),
        puis les locaux. Retient le premier dans l'ordre d'apparition dans le texte.
        Validation : minimum 7 chiffres, rejet des plages d'années.
        """
        candidates = []

        for m in self.PHONE_INTL.finditer(text):
            candidate = m.group(0).strip()
            digits    = re.sub(r"\D", "", candidate)
            if len(digits) >= 7 and not self._YEAR_RANGE.match(candidate):
                candidates.append((m.start(), candidate))

        for m in self.PHONE_LOCAL.finditer(text):
            candidate = m.group(0).strip()
            if not self._YEAR_RANGE.match(candidate):
                candidates.append((m.start(), candidate))

        if not candidates:
            return None

        candidates.sort(key=lambda x: x[0])
        return candidates[0][1]


# ---------------------------------------------------------------------------
# COUCHE 2 — NER spaCy (entités nommées)
# ---------------------------------------------------------------------------

# Tokens qui, s'ils apparaissent sur une ligne, l'excluent comme nom de candidat
_NAME_LINE_BLACKLIST: set[str] = {
    # Titres de poste (FR + EN)
    "developer", "engineer", "manager", "analyst", "architect", "consultant",
    "designer", "intern", "trainee", "officer", "lead", "senior", "junior",
    "stagiaire", "ingénieur", "développeur", "responsable", "chef", "directeur",
    # Rubriques CV
    "cv", "resume", "curriculum", "vitae", "profil", "profile", "portfolio",
    "about", "contact", "summary", "formation", "compétences", "skills",
    "experience", "expérience", "education", "projet", "projects",
    "certifications", "langues", "languages", "références", "references",
    # Technologies
    "python", "java", "javascript", "typescript", "react", "angular", "vue",
    "docker", "kubernetes", "aws", "azure", "gcp", "sql", "linux", "git",
    "databricks", "spark", "kafka", "hadoop", "tensorflow", "pytorch",
    # Marqueurs de contact
    "email", "phone", "tel", "mobile", "adresse", "address", "linkedin", "github",
}

# Mot de nom valide : lettres (dont accents), tirets, apostrophes — pas de chiffres
_NAME_WORD_RE = re.compile(r"^[A-Za-zÀ-ÿ][A-Za-z\u00C0-\u024F'\-]*$")

# Exclut les lignes entièrement en majuscules (rubriques : "EXPÉRIENCE PROFESSIONNELLE")
_ALL_CAPS_RE  = re.compile(r"^[A-ZÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝ\s\-]+$")


def _is_valid_name_line(line: str) -> bool:
    """
    Retourne True si la ligne est un candidat plausible pour un nom de candidat.

    Filtres :
        1. Longueur entre 3 et 60 chars
        2. Pas d'email, URL, date, bullet point
        3. Pas entièrement en majuscules (= rubrique)
        4. Entre 2 et 6 mots, tous composés de lettres/tirets/apostrophes
        5. Au moins 2 mots avec initiale majuscule
        6. Aucun mot blacklisté
    """
    stripped = line.strip()
    if not stripped or len(stripped) < 3 or len(stripped) > 60:
        return False
    if re.search(r"@", stripped):
        return False
    if re.search(r"https?://|www\.|linkedin\.com|github\.com", stripped):
        return False
    if re.search(r"\b(19|20)\d{2}\b", stripped):
        return False
    if re.match(r"^[\•\-\*\▸\◦\→]", stripped):
        return False
    if _ALL_CAPS_RE.match(stripped):
        return False

    clean = re.sub(r"[|·—/]", " ", stripped)
    words = [w for w in clean.split() if w]

    if len(words) < 2 or len(words) > 6:
        return False
    if not all(_NAME_WORD_RE.match(w) for w in words):
        return False
    if sum(1 for w in words if w[0].isupper()) < 2:
        return False
    if {w.lower() for w in words} & _NAME_LINE_BLACKLIST:
        return False

    return True


def _extract_name_from_header(text: str, max_lines: int = 10) -> str | None:
    """
    Scanne les premières lignes du CV pour extraire le nom du candidat.

    Heuristique : le nom est presque toujours dans les 10 premières lignes non
    vides, seul sur sa ligne, sans email/URL/chiffres, avec ≥2 mots capitalisés.
    Plus robuste que spaCy pour les noms maghrébins et arabes peu représentés
    dans les corpus d'entraînement.
    """
    lines     = text.split("\n")
    non_empty = [l for l in lines if l.strip()][:max_lines]
    for line in non_empty:
        if _is_valid_name_line(line):
            return line.strip()
    return None


class NERExtractor:
    """
    Extraction des entités nommées via spaCy + heuristique d'en-tête.

    Cibles :
        - full_name → heuristique lignes d'abord, spaCy (PER) en fallback
        - location  → spaCy GPE/LOC avec fusion ville + pays contigus
        - ORG       → entreprises et établissements (pour validation LLM)

    Pourquoi l'heuristique en premier pour le nom :
        spaCy fr_core_news_sm reconnaît
        mal les noms maghrébins et tunisiens. L'heuristique basée sur la
        structure du CV est plus fiable sur ces profils.
    """

    _MODEL_MAP = {
        "fr": ["fr_core_news_sm", "fr_core_news_md", "fr_core_news_lg"],
        "en": ["en_core_web_sm",  "en_core_web_md",  "en_core_web_lg"],
    }

    _ORG_BLACKLIST = {
        "github", "linkedin", "microsoft", "google", "apple",
        "python", "javascript", "typescript", "react", "docker",
    }

    _NER_NAME_BLACKLIST = {
        "azure databricks", "google cloud", "amazon web services",
        "microsoft azure", "apache kafka", "apache spark",
        "blog personnel", "portfolio", "personal website",
        "curriculum vitae", "resume", "profil", "profile",
        "backend developer", "frontend developer", "full stack developer",
        "software engineer", "data scientist", "data engineer",
        "devops engineer", "product manager", "project manager",
        "lead developer", "senior developer", "junior developer",
    }

    def __init__(self):
        self.models: dict[str, spacy.language.Language] = {}

    def _get_nlp(self, lang: str) -> spacy.language.Language:
        """Charge et retourne le modèle spaCy pour la langue détectée, avec cache."""
        if lang in self.models:
            return self.models[lang]

        model_list = self._MODEL_MAP.get(lang, self._MODEL_MAP["fr"])

        for model_name in model_list:
            try:
                nlp = spacy.load(model_name)
                logger.info("Modèle spaCy chargé pour %s : %s", lang, model_name)
                self.models[lang] = nlp
                return nlp
            except OSError:
                continue

        logger.warning(
            "Aucun modèle spaCy pour '%s'. NER désactivé.\n"
            "  → python -m spacy download %s", lang, model_list[0]
        )
        nlp = spacy.blank(lang if lang in ("fr", "en") else "fr")
        self.models[lang] = nlp
        return nlp

    def extract(self, text: str) -> dict:
        """
        Extrait les entités nommées pertinentes du texte CV.

        Stratégie full_name :
            1. Heuristique d'en-tête — fiable pour noms maghrébins et CVs classiques
            2. spaCy NER (PER) — fallback si l'heuristique ne trouve rien
        """
        # Détection de la langue
        try:
            lang = detect(text[:5000])
            lang = lang if lang in self._MODEL_MAP else "fr"
            logger.info("Langue détectée : %s", lang)
        except Exception as e:
            logger.warning("Échec détection langue : %s → fallback fr", e)
            lang = "fr"

        header_text = text[:10_000]

        # Étape 1 : heuristique d'en-tête (prioritaire)
        heuristic_name = _extract_name_from_header(text)

        # Étape 2 : NER spaCy
        nlp = self._get_nlp(lang)
        doc = nlp(header_text)

        ner_persons   = []
        location_ents = []
        organizations = []

        for ent in doc.ents:
            val   = ent.text.strip()
            label = ent.label_

            if not val or len(val) < 2:
                continue

            if label in ("PER", "PERSON"):
                words = val.split()
                if (len(words) >= 2
                        and all(w[0].isupper() for w in words if w)
                        and val.lower() not in self._NER_NAME_BLACKLIST):
                    ner_persons.append(val)

            elif label in ("GPE", "LOC", "LOCATION"):
                location_ents.append((ent.start_char, val))

            elif label in ("ORG", "ORGANIZATION"):
                if val.lower() not in self._ORG_BLACKLIST:
                    organizations.append(val)

        # Fusion : heuristique prioritaire, spaCy en fallback
        full_name = heuristic_name or (ner_persons[0] if ner_persons else None)
        location  = self._merge_location(location_ents, header_text)

        result = {
            "full_name":      full_name,
            "location":       location,
            "_organizations": list(dict.fromkeys(organizations)),
        }

        logger.info(
            "NER → full_name=%s (src=%s) | location=%s | orgs=%d",
            result["full_name"],
            "heuristic" if heuristic_name else "spacy",
            result["location"],
            len(result["_organizations"]),
        )
        return result

    @staticmethod
    def _merge_location(location_ents: list[tuple], source_text: str) -> str | None:
        """
        Fusionne des entités GPE consécutives en une localisation complète.
        Ex : "Sfax" + "Tunisie" à moins de 20 chars d'écart → "Sfax, Tunisie"
        """
        if not location_ents:
            return None
        if len(location_ents) == 1:
            return location_ents[0][1]

        first_pos, first_val = location_ents[0]
        window               = source_text[first_pos + len(first_val) : first_pos + len(first_val) + 20]
        _, second_val        = location_ents[1]

        if second_val in window:
            return f"{first_val}, {second_val}"
        return first_val


# ---------------------------------------------------------------------------
# COUCHE 3 — LLM Groq (sémantique)
# ---------------------------------------------------------------------------

class LLMExtractor:
    """
    Extraction sémantique via LLM Groq.

    Responsabilités :
        - Résumé professionnel
        - Compétences techniques (avec inférence contextuelle)
        - Expériences structurées (titre, entreprise, période, missions)
        - Formation
        - Certifications
        - Langues avec niveaux

    Le LLM reçoit les champs déjà extraits par Regex + NER comme contexte,
    ce qui réduit les hallucinations sur les champs déjà connus.

    Modèles production Groq (131 072 tokens de contexte) :
        llama-3.1-8b-instant        560 t/s   $0.05/$0.08   volume, rapidité
        llama-3.3-70b-versatile     280 t/s   $0.59/$0.79   qualité production
        openai/gpt-oss-20b         1000 t/s   $0.075/$0.30  vitesse maximale
        openai/gpt-oss-120b         500 t/s   $0.15/$0.60   meilleure qualité

    Modèles preview (évaluation uniquement) :
        meta-llama/llama-4-scout-17b-16e-instruct   750 t/s
        qwen/qwen3-32b                              400 t/s  (multilingue FR/AR/EN)
        moonshotai/kimi-k2-instruct-0905            200 t/s  (contexte 262k)
    """

    _CHAR_LIMIT = 60_000  # ~800 tokens réservés pour le prompt, marge sur 131k contexte

    _JSON_SCHEMA = """{
  "full_name":      "string | null",
  "location":       "string | null  (ville, pays)",
  "summary":        "string | null  (2-3 phrases synthétiques)",
  "skills":         ["string"],
  "languages":      [{"language": "string", "level": "string"}],
  "experiences": [
    {
      "title":       "string",
      "company":     "string",
      "location":    "string | null",
      "period":      "string  (ex: Jan 2022 - Présent)",
      "description": "string  (missions principales)"
    }
  ],
  "education": [
    {
      "degree":      "string",
      "institution": "string",
      "location":    "string | null",
      "period":      "string"
    }
  ],
  "certifications": ["string"]
}"""

    def __init__(self, model: str = "llama-3.1-8b-instant", temperature: float = 0.0):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise EnvironmentError(
                "GROQ_API_KEY manquant. Créer un fichier .env avec GROQ_API_KEY=..."
            )
        self.client      = Groq(api_key=api_key)
        self.model       = model
        self.temperature = temperature

    def extract(self, cv_text: str, pre_extracted: dict) -> dict:
        """
        Lance l'extraction LLM en fournissant le contexte pré-extrait.

        Args:
            cv_text:       Texte brut du CV
            pre_extracted: Champs déjà extraits (Regex + NER)

        Returns:
            dict structuré selon le schéma JSON
        """
        # Troncature intelligente : on coupe au dernier saut de ligne
        if len(cv_text) > self._CHAR_LIMIT:
            truncated = cv_text[:self._CHAR_LIMIT].rsplit("\n", 1)[0]
            logger.warning(
                "CV tronqué : %d → %d chars (limite %d)",
                len(cv_text), len(truncated), self._CHAR_LIMIT,
            )
        else:
            truncated = cv_text

        # Injection du contexte pré-extrait pour ancrer le LLM
        known = {k: v for k, v in pre_extracted.items() if v and not k.startswith("_")}
        context_block = (
            f"\nINFORMATIONS DÉJÀ EXTRAITES (ne pas modifier) :\n"
            f"{json.dumps(known, ensure_ascii=False, indent=2)}\n"
            if known else ""
        )

        prompt = f"""Tu es un expert en analyse de CV et recrutement.

Analyse le CV ci-dessous et extrais les informations manquantes en respectant le schéma JSON.
{context_block}
RÈGLES :
- Réponds UNIQUEMENT avec du JSON valide, sans markdown, sans explication.
- Si une information est absente, utilise null ou [].
- Ne complète pas, n'invente pas d'informations.
- Pour les compétences, liste uniquement ce qui est explicitement mentionné.
- Normalise les dates au format "MMM YYYY" (ex: Jan 2022).
- Dans les valeurs string, remplace tout guillemet " par une apostrophe ' pour éviter de casser le JSON.

SCHÉMA JSON :
{self._JSON_SCHEMA}

CV :
---
{truncated}
---

JSON :"""

        raw = self.client.chat.completions.create(
            model       = self.model,
            messages    = [{"role": "user", "content": prompt}],
            temperature = self.temperature,
        ).choices[0].message.content

        return self._parse_json_response(raw)

 #Fonction de génération de résumé profil par IA si absent dans le CV       

    def generate_summary(self, cv_text: str, extracted: dict) -> str:
        """
        Demande au LLM de rédiger un résumé professionnel du candidat
        à partir du texte brut et des données déjà extraites.

        Appelée uniquement si le CV ne contient pas de résumé existant.
        Retourne une chaîne de 2-3 phrases maximum.
        """
        # Construire un contexte compact pour le LLM
        context_parts = []

        if extracted.get("full_name"):
            context_parts.append(f"Candidat : {extracted['full_name']}")

        if extracted.get("skills"):
            skills = extracted["skills"]
            if isinstance(skills, list):
                context_parts.append(f"Compétences : {', '.join(skills[:15])}")

        if extracted.get("experiences"):
            exps = extracted["experiences"]
            if isinstance(exps, list) and exps:
                titles = [e.get("title", "") for e in exps[:3] if e.get("title")]
                if titles:
                    context_parts.append(f"Expériences récentes : {', '.join(titles)}")

        if extracted.get("education"):
            edus = extracted["education"]
            if isinstance(edus, list) and edus:
                deg = edus[0].get("degree", "")
                inst = edus[0].get("institution", "")
                if deg:
                    context_parts.append(f"Formation : {deg}{f' — {inst}' if inst else ''}")

        if extracted.get("languages"):
            langs = extracted["languages"]
            if isinstance(langs, list):
                lang_str = ", ".join(
                    f"{l.get('language', '')} ({l.get('level', '')})"
                    for l in langs[:4] if l.get("language")
                )
                if lang_str:
                    context_parts.append(f"Langues : {lang_str}")

        context_block = "\n".join(context_parts)

        # Limiter le texte brut pour ne pas dépasser le contexte
        cv_truncated = cv_text[:8000]

        prompt = f"""Tu es un expert en recrutement. Rédige un résumé professionnel concis (2-3 phrases maximum) pour ce candidat.

    Le résumé doit :
    - Mentionner le domaine d'expertise principal
    - Citer 2-3 compétences clés
    - Indiquer le niveau d'expérience si possible
    - Être rédigé à la 3e personne (ex: "Développeur Full Stack avec...")
    - Être en français

    Données extraites :
    {context_block}

    Extrait du CV :
    ---
    {cv_truncated}
    ---

    Réponds UNIQUEMENT avec le texte du résumé, sans guillemets, sans introduction, sans explication."""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,  # légèrement créatif mais cohérent
                max_tokens=150,    # 2-3 phrases max
            )
            summary = response.choices[0].message.content.strip()

            # Nettoyage : enlever les guillemets si le LLM en ajoute
            summary = summary.strip('"\'')

            logger.info("Résumé IA généré : %d caractères", len(summary))
            return summary

        except Exception as e:
            logger.warning("Échec génération résumé IA : %s", e)
            return ""




    @staticmethod
    def _parse_json_response(raw: str) -> dict:
        """Nettoie et parse la réponse JSON du LLM.

        Stratégie de récupération en cascade :
            1. json.loads standard  — cas nominal
            2. json_repair          — corrige les guillemets non échappés,
                                      virgules manquantes, JSON tronqué
                                      (pip install json-repair)
            3. Extraction regex     — dernier recours sur champs scalaires
        """
        # Nettoyage des balises markdown
        cleaned = re.sub(r"```(?:json)?", "", raw).replace("```", "").strip()
        start   = cleaned.find("{")
        end     = cleaned.rfind("}") + 1

        if start == -1 or end == 0:
            logger.error("Aucun JSON dans la réponse LLM : %s", raw[:200])
            return {}

        json_str = cleaned[start:end]

        # Tentative 1 : parsing standard
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            pass

        # Tentative 2 : json_repair — gère les guillemets non échappés,
        # virgules manquantes, JSON tronqué (réponse coupée par max_tokens)
        try:
            from json_repair import repair_json
            repaired = repair_json(json_str)
            result   = json.loads(repaired)
            logger.warning("JSON réparé via json_repair (réponse LLM malformée)")
            return result
        except ImportError:
            logger.debug("json_repair non installé. pip install json-repair")
        except (json.JSONDecodeError, Exception):
            pass

        # Tentative 3 : extraction regex des champs scalaires — dernier recours
        logger.error(
            "JSON irrécupérable — extraction partielle des champs scalaires.\n"
            "Réponse LLM (300 premiers chars) : %s", json_str[:300]
        )
        return LLMExtractor._extract_scalars_fallback(json_str)

    @staticmethod
    def _extract_scalars_fallback(text: str) -> dict:
        """Extrait les champs scalaires par regex quand le JSON est irrécupérable.

        Ne récupère que les champs simples (strings et null) — les listes
        sont trop complexes à extraire par regex et valent mieux être laissées
        vides que mal parsées.
        """
        result: dict = {
            "skills": [], "languages": [], "experiences": [],
            "education": [], "certifications": [],
        }
        scalar_fields = ("full_name", "location", "summary", "email",
                         "phone", "linkedin", "github")

        for field_name in scalar_fields:
            # Cherche : "field_name": "valeur" ou "field_name": null
            pattern = rf'"{field_name}"\s*:\s*(?:"((?:[^"\\]|\\.)*)"|(null))'
            match   = re.search(pattern, text)
            if match:
                result[field_name] = match.group(1)  # None si c'était null
            else:
                result[field_name] = None

        return result


# ---------------------------------------------------------------------------
# COUCHE 4 — Post-processing & Validation croisée
# ---------------------------------------------------------------------------

class PostProcessor:
    """
    Fusion, validation croisée et normalisation des résultats des 3 couches.

    Règles de priorité par champ :
        email, phone, linkedin, github  → Regex > LLM  (validation croisée si LLM)
        full_name, location             → NER   > LLM
        summary, skills, exp, edu, lang → LLM uniquement (sémantique pure)
    """

    _EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")

    def merge(
        self,
        regex_data: dict,
        ner_data:   dict,
        llm_data:   dict,
        raw_text:   str,
    ) -> dict:
        """
        Fusionne les résultats des 3 couches avec règles de priorité.

        Returns:
            dict final normalisé
        """
        merged = {}

        # Champs déterministes : Regex prime, LLM validé en fallback
        for f in ("email", "phone", "linkedin", "github"):
            regex_val = regex_data.get(f)
            llm_val   = llm_data.get(f)

            if regex_val:
                merged[f] = regex_val
            elif llm_val:
                merged[f] = self._validate_against_source(llm_val, raw_text, f)
            else:
                merged[f] = None

        # Entités nommées : NER prioritaire, LLM en fallback
        merged["full_name"] = ner_data.get("full_name") or llm_data.get("full_name")
        merged["location"]  = ner_data.get("location")  or llm_data.get("location")

        # Champs sémantiques : LLM uniquement
        for f in ("summary", "skills", "languages", "experiences", "education", "certifications"):
            merged[f] = llm_data.get(f) or ([] if f != "summary" else None)

        # Normalisation finale
        merged = self._normalize(merged)

        # Métadonnées de debug
        merged["_extraction_meta"] = {
            "regex_found":    {k: bool(v) for k, v in regex_data.items()},
            "ner_found":      {k: bool(v) for k, v in ner_data.items() if not k.startswith("_")},
            "llm_overridden": self._compute_overrides(regex_data, ner_data, llm_data),
        }

        return merged

    def _validate_against_source(self, llm_value: str, raw_text: str, field: str) -> Optional[str]:
        """
        Vérifie qu'une valeur extraite par le LLM est bien présente dans le texte source.
        Protège contre les hallucinations sur les champs critiques.
        """
        if not llm_value:
            return None

        if field == "email":
            if llm_value.lower() in raw_text.lower() and self._EMAIL_RE.match(llm_value):
                return llm_value
            logger.warning("Email LLM '%s' introuvable dans la source → invalide", llm_value)
            return None

        # Autres champs : vérification souple (on garde avec avertissement)
        if llm_value.lower() not in raw_text.lower():
            logger.debug("Valeur LLM '%s' non vérifiable dans la source (%s)", llm_value, field)
        return llm_value

    @staticmethod
    def _normalize(data: dict) -> dict:
        """Normalisation finale : types, casse, déduplication."""
        # Strings vides → None
        for f in ("full_name", "email", "phone", "location", "linkedin", "github", "summary"):
            val = data.get(f)
            if isinstance(val, str) and not val.strip():
                data[f] = None

        # Skills : lowercase, déduplication, tri
        if data.get("skills"):
            data["skills"] = sorted({
                s.strip().lower()
                for s in data["skills"]
                if isinstance(s, str) and s.strip()
            })

        # Certifications : déduplication (ordre préservé)
        if data.get("certifications"):
            data["certifications"] = list(dict.fromkeys(
                c.strip() for c in data["certifications"]
                if isinstance(c, str) and c.strip()
            ))

        # Listes → liste vide si None
        for f in ("skills", "languages", "experiences", "education", "certifications"):
            if not isinstance(data.get(f), list):
                data[f] = []

        return data

    @staticmethod
    def _compute_overrides(regex_data: dict, ner_data: dict, llm_data: dict) -> list[str]:
        """Retourne les champs où Regex/NER a écrasé le résultat LLM."""
        overrides = []
        for f in ("email", "phone", "linkedin", "github"):
            if regex_data.get(f) and llm_data.get(f) and regex_data[f] != llm_data.get(f):
                overrides.append(f)
        if (ner_data.get("full_name") and llm_data.get("full_name")
                and ner_data["full_name"] != llm_data.get("full_name")):
            overrides.append("full_name")
        return overrides


# ---------------------------------------------------------------------------
# Pipeline Principal
# ---------------------------------------------------------------------------

class CVParser:
    """
    Orchestrateur du pipeline d'extraction en 4 couches.

    Usage :
        parser = CVParser()
        result = parser.parse("cv.pdf")
        print(result.to_json())

    Usage avec modèle spécifique :
        parser = CVParser(llm_model="llama-3.3-70b-versatile")
        result = parser.parse("cv_scanné.png")
    """

    def __init__(
        self,
        llm_model:       str   = "llama-3.1-8b-instant",
        llm_temperature: float = 0.0,
    ):
        self.text_extractor  = TextExtractor()
        self.regex_extractor = RegexExtractor()
        self.ner_extractor   = NERExtractor()
        self.llm_extractor   = LLMExtractor(model=llm_model, temperature=llm_temperature)
        self.post_processor  = PostProcessor()

        logger.info("CVParser initialisé — LLM : %s", llm_model)

    def parse(self, file_path: str) -> CVData:
        """
        Parse un CV depuis un fichier (PDF, DOCX, TXT, image).

        Returns:
            CVData structuré et normalisé.
        """
        logger.info("═══ Parsing : %s ═══", Path(file_path).name)

        # Couche 0 : extraction du texte brut
        raw_text, extraction_method = TextExtractor.extract(file_path)
        logger.info(
            "Extraction texte : méthode=%s, longueur=%d chars",
            extraction_method, len(raw_text),
        )

        if len(raw_text) < 20:
            raise ValueError(
                f"Texte extrait insuffisant ({len(raw_text)} chars). "
                f"Le fichier est peut-être corrompu ou vide."
            )

        return self.parse_text(raw_text)

    def parse_text(self, raw_text: str) -> CVData:
        """
        Parse un CV depuis du texte brut déjà extrait.
        Utile pour les benchmarks et les tests unitaires.
        """
        # Couche 1 : Regex
        logger.info("Couche 1 — Regex")
        regex_data = self.regex_extractor.extract(raw_text)

        # Couche 2 : NER spaCy
        logger.info("Couche 2 — NER spaCy")
        ner_data = self.ner_extractor.extract(raw_text)

        # Couche 3 : LLM Groq
        logger.info("Couche 3 — LLM Groq")
        pre_extracted = {
            **regex_data,
            **{k: v for k, v in ner_data.items() if not k.startswith("_")},
        }
        llm_data = self.llm_extractor.extract(raw_text, pre_extracted)

        # Couche 4 : Post-processing & fusion
        logger.info("Couche 4 — Post-processing & fusion")
        merged = self.post_processor.merge(regex_data, ner_data, llm_data, raw_text)

        # ── Génération du résumé IA si absent ────────────────────────────
        # Le CV peut ne pas contenir de section "résumé" ou "profil".
        # Dans ce cas on demande au LLM de le rédiger à partir des données.
        if not merged.get("summary"):
            logger.info("Résumé absent du CV → génération par le LLM")
            merged["summary"] = self.llm_extractor.generate_summary(raw_text, merged)
        else:
            logger.info("Résumé existant conservé (%d chars)", len(merged["summary"]))

        return CVData.from_dict(merged)


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------

class CVParserFactory:
    """
    Presets préconfigurés selon le cas d'usage.

    Modèles production Groq (131 072 tokens, mars 2026) :
        fast()       llama-3.1-8b-instant        560 t/s   $0.05/$0.08
        accurate()   llama-3.3-70b-versatile     280 t/s   $0.59/$0.79
        fastest()    openai/gpt-oss-20b          1000 t/s  $0.075/$0.30
        best()       openai/gpt-oss-120b          500 t/s  $0.15/$0.60

    Modèles preview (évaluation uniquement) :
        multilingual()  qwen/qwen3-32b — CVs FR/AR/EN mixtes
    """

    @staticmethod
    def fast() -> CVParser:
        """Volume et rapidité — 95% des CVs standards. llama-3.1-8b-instant."""
        return CVParser(llm_model="llama-3.1-8b-instant")

    @staticmethod
    def accurate() -> CVParser:
        """Qualité production — CVs complexes. llama-3.3-70b-versatile."""
        return CVParser(llm_model="llama-3.3-70b-versatile")

    @staticmethod
    def fastest() -> CVParser:
        """Vitesse absolue — APIs temps réel. openai/gpt-oss-20b."""
        return CVParser(llm_model="openai/gpt-oss-20b")

    @staticmethod
    def best() -> CVParser:
        """Meilleure qualité — CVs très complexes. openai/gpt-oss-120b."""
        return CVParser(llm_model="openai/gpt-oss-120b")

    @staticmethod
    def multilingual() -> CVParser:
        """CVs FR/AR/EN mixtes. qwen/qwen3-32b (preview)."""
        return CVParser(llm_model="qwen/qwen3-32b")

    @staticmethod
    def custom(model: str) -> CVParser:
        """Modèle personnalisé — pour tests et expérimentations."""
        return CVParser(llm_model=model)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage : python cv_parser.py <chemin_cv> [modele]")
        print()
        print("Exemples :")
        print("  python cv_parser.py cv.pdf")
        print("  python cv_parser.py cv_scanné.jpg")
        print("  python cv_parser.py cv.docx llama-3.3-70b-versatile")
        print()
        print("Formats supportés : pdf, docx, txt, md, jpg, png, webp, tiff")
        sys.exit(1)

    path  = sys.argv[1]
    model = sys.argv[2] if len(sys.argv) > 2 else None

    parser = CVParserFactory.custom(model) if model else CVParserFactory.fast()
    result = parser.parse(path)

    print(result.to_json(include_meta=False))